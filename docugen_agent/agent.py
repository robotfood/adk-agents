import datetime
from zoneinfo import ZoneInfo
from google.adk.agents import Agent
# from google.adk.agents.llm_output_formatters import SimpleFormatter

from docx import Document
from docx.shared import Inches

def create_sample_doc_with_text(llm_response: str):
    # Create a new Document object
    document = Document()

    # Add a heading
    document.add_heading('My Sample Document', level=1)

    # Add a simple paragraph
    document.add_paragraph(llm_response)

    filename = "sample_document.docx"

    # Save the document
    document.save(filename)
    print(f"'{filename}' created successfully with text content.")
    return {
        "status": "success",
        "report": (
            "Successfully created a document with the provided text content."
        ),
    }

root_agent = Agent(
    name="document_generation_agent",  # Fixed typo in name
    model="gemini-2.0-flash",
    description="An agent to create documents from LLM queries.",
    instruction=(
        "You are a helpful agent who can answer user questions, and if needed automatically create formatted documents from your responses based on user input."
    ),
    tools=[create_sample_doc_with_text]
)
